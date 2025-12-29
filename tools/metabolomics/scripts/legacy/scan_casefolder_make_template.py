#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import json
import argparse
from datetime import datetime
from pathlib import Path

import pandas as pd

# ---- optional deps (best-effort) ----
PDFPLUMBER_OK = True
try:
    import pdfplumber
except Exception:
    PDFPLUMBER_OK = False

PYMUPDF_OK = True
try:
    import fitz  # PyMuPDF
except Exception:
    PYMUPDF_OK = False

DOCX_OK = True
try:
    import docx  # python-docx
except Exception:
    DOCX_OK = False

PIL_OK = True
try:
    from PIL import Image
except Exception:
    PIL_OK = False

OPENPYXL_OK = True
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    from openpyxl.worksheet.datavalidation import DataValidation
except Exception:
    OPENPYXL_OK = False


# -----------------------------
# helpers
# -----------------------------
def ensure_dir(p: Path) -> Path:
    p.mkdir(parents=True, exist_ok=True)
    return p

def guess_sample_id_from_name(fname: str) -> str:
    # 你可以按你们规则改：比如 “NC1/HS3/ZS2...” 等
    stem = Path(fname).stem
    stem = re.sub(r"\s+", "", stem)
    return stem

def normalize_text(s: str) -> str:
    if not s:
        return ""
    s = s.replace("\u3000", " ")
    s = re.sub(r"[ \t]+", " ", s)
    return s

def try_parse_datetime(s: str):
    """
    尝试从字符串里抓一个 datetime（失败就返回 None）
    支持常见：YYYY-MM-DD HH:MM / YYYY/MM/DD HH:MM / YYYY年MM月DD日HH时MM分 等
    """
    if not s:
        return None

    s0 = s
    # 统一分隔符
    s = s.replace("年", "-").replace("月", "-").replace("日", " ").replace("时", ":").replace("分", "")
    s = s.replace("/", "-")
    s = re.sub(r"\s+", " ", s).strip()

    fmts = [
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d",
    ]
    # 先抓一个像日期时间的片段
    m = re.search(r"(\d{4}-\d{1,2}-\d{1,2}(?:\s+\d{1,2}:\d{2}(?::\d{2})?)?)", s)
    cand = m.group(1) if m else s

    for f in fmts:
        try:
            return datetime.strptime(cand, f)
        except Exception:
            pass

    # 兜底：可能是 “12-28 15:30” 没年份（不建议自动补）
    return None

def extract_fields_from_text(text: str) -> dict:
    """
    从“可复制的文本”中尽力提取：
    NIHSS、发病时间/到院时间/采样时间(原文片段)、年龄、性别、HTN、DM、溶栓/取栓、用药(粗抓)
    """
    t = normalize_text(text)

    out = {}

    # NIHSS
    m = re.search(r"NIHSS\s*[:：]?\s*(\d{1,2})", t, flags=re.IGNORECASE)
    if m:
        out["NIHSS"] = int(m.group(1))

    # 年龄
    m = re.search(r"(?:年龄|Age)\s*[:：]?\s*(\d{1,3})", t, flags=re.IGNORECASE)
    if m:
        out["Age"] = int(m.group(1))

    # 性别
    # 注意：很多病历写“男/女”或“性别:男”
    m = re.search(r"(?:性别|Sex)\s*[:：]?\s*(男|女|M|F)", t, flags=re.IGNORECASE)
    if m:
        v = m.group(1).upper()
        out["Sex"] = "M" if v in ["男", "M"] else "F"

    # HTN / DM（粗判定：出现即“是”，不出现留空）
    if re.search(r"(高血压|HTN|Hypertension)", t, flags=re.IGNORECASE):
        out["HTN"] = "Yes"
    if re.search(r"(糖尿病|DM|Diabetes)", t, flags=re.IGNORECASE):
        out["DM"] = "Yes"

    # 溶栓 / 取栓
    if re.search(r"(溶栓|rt-PA|tPA|阿替普酶)", t, flags=re.IGNORECASE):
        out["Thrombolysis"] = "Yes"
    if re.search(r"(取栓|机械取栓|thrombectomy)", t, flags=re.IGNORECASE):
        out["Thrombectomy"] = "Yes"

    # 发病/起病时间、到院时间、采样/采集时间：抓“标签+后面一段”
    def grab_line(keywords, fieldname):
        for kw in keywords:
            m = re.search(rf"{kw}\s*[:：]?\s*([^\n\r;，。]{{6,40}})", t)
            if m:
                out[fieldname] = m.group(1).strip()
                dt = try_parse_datetime(out[fieldname])
                if dt:
                    out[fieldname + "_parsed"] = dt.strftime("%Y-%m-%d %H:%M")
                return

    grab_line(["发病时间", "起病时间", "Onset time"], "OnsetTime")
    grab_line(["首次到院时间", "到院时间", "Arrival time"], "FirstArrivalTime")
    grab_line(["采样时间", "采集时间", "Sampling time"], "SamplingTime")

    # 发病小时数：如果文本中直接写了“发病xx小时”
    m = re.search(r"(?:发病|起病)\s*(\d+(?:\.\d+)?)\s*(?:小时|h)", t)
    if m:
        out["OnsetHours"] = float(m.group(1))

    # 用药（非常粗：抓“用药:”后 10~80 字）
    m = re.search(r"(?:用药|Medications?)\s*[:：]?\s*([^\n\r]{10,80})", t, flags=re.IGNORECASE)
    if m:
        out["Meds_raw"] = m.group(1).strip()

    return out

def pdf_has_extractable_text(pdf_path: Path, max_pages=2) -> bool:
    if not PDFPLUMBER_OK:
        return False
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            for i, page in enumerate(pdf.pages[:max_pages]):
                txt = page.extract_text() or ""
                if len(txt.strip()) >= 50:
                    return True
    except Exception:
        return False
    return False

def extract_text_from_pdf(pdf_path: Path, max_pages=3) -> str:
    if not PDFPLUMBER_OK:
        return ""
    texts = []
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages[:max_pages]:
                texts.append(page.extract_text() or "")
    except Exception:
        return ""
    return "\n".join(texts)

def render_pdf_preview(pdf_path: Path, out_png: Path, zoom=2.0) -> bool:
    """
    用 PyMuPDF 把 PDF 第1页渲染成 PNG，给人工快速浏览
    """
    if not PYMUPDF_OK:
        return False
    try:
        doc = fitz.open(str(pdf_path))
        page = doc.load_page(0)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        pix.save(str(out_png))
        doc.close()
        return True
    except Exception:
        return False

def extract_text_from_docx(docx_path: Path) -> str:
    if not DOCX_OK:
        return ""
    try:
        d = docx.Document(str(docx_path))
        parts = []
        for p in d.paragraphs:
            if p.text:
                parts.append(p.text)
        for tb in d.tables:
            for row in tb.rows:
                parts.append(" ".join([c.text for c in row.cells if c.text]))
        return "\n".join(parts)
    except Exception:
        return ""

def image_to_preview(img_path: Path, out_png: Path) -> bool:
    """
    如果是图片就复制/转 png，便于 Excel 超链接
    """
    try:
        if not PIL_OK:
            # 没 PIL 就直接复制
            out_png.write_bytes(img_path.read_bytes())
            return True
        im = Image.open(str(img_path))
        im = im.convert("RGB")
        im.save(str(out_png))
        return True
    except Exception:
        return False

def make_excel_template(rows, out_xlsx: Path):
    if not OPENPYXL_OK:
        # fallback csv
        pd.DataFrame(rows).to_csv(out_xlsx.with_suffix(".csv"), index=False, encoding="utf-8-sig")
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "ClinicalEntry"

    headers = [
        "SampleID","Group(HS/ZS/NC)","SourceFile","PreviewLink",
        "SamplingTime","OnsetTime","FirstArrivalTime","OnsetHours",
        "NIHSS","Age","Sex(M/F)","HTN(Yes/No)","DM(Yes/No)",
        "Meds","Thrombectomy(Yes/No)","Thrombolysis(Yes/No)",
        "Notes"
    ]
    ws.append(headers)

    # style
    bold = Font(bold=True)
    for j, h in enumerate(headers, start=1):
        c = ws.cell(row=1, column=j)
        c.font = bold
        c.alignment = Alignment(horizontal="center", vertical="center")

    # data rows
    for r in rows:
        ws.append([r.get(h, "") for h in headers])

    # column width
    widths = {
        "A":16, "B":14, "C":40, "D":18, "E":18, "F":18, "G":18, "H":12,
        "I":8, "J":6, "K":9, "L":10, "M":10, "N":30, "O":18, "P":18, "Q":30
    }
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    # validations
    dv_group = DataValidation(type="list", formula1='"HS,ZS,NC"', allow_blank=True)
    dv_yesno = DataValidation(type="list", formula1='"Yes,No"', allow_blank=True)
    dv_sex = DataValidation(type="list", formula1='"M,F"', allow_blank=True)

    ws.add_data_validation(dv_group)
    ws.add_data_validation(dv_yesno)
    ws.add_data_validation(dv_sex)

    n = len(rows) + 1
    dv_group.add(f"B2:B{n}")
    dv_sex.add(f"K2:K{n}")
    dv_yesno.add(f"L2:L{n}")
    dv_yesno.add(f"M2:M{n}")
    dv_yesno.add(f"O2:O{n}")
    dv_yesno.add(f"P2:P{n}")

    # freeze
    ws.freeze_panes = "A2"

    # instructions sheet
    ws2 = wb.create_sheet("README")
    ws2["A1"] = "How to use"
    ws2["A1"].font = bold
    ws2["A2"] = "1) Open ClinicalEntry sheet"
    ws2["A3"] = "2) Click PreviewLink to view the first page preview"
    ws2["A4"] = "3) Fill key clinical variables. Auto-extracted fields are only suggestions—please verify."
    ws2.column_dimensions["A"].width = 100

    wb.save(str(out_xlsx))

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--indir", required=True, help="病历文件夹 (WSL路径), e.g. /mnt/h/Data/Yuchun-yanshi/眼屎收集病历")
    ap.add_argument("--outdir", default="Clinical_Extract_Out", help="输出目录")
    ap.add_argument("--max_pages_text", type=int, default=3, help="PDF抽取文本的最大页数（可复制文本PDF）")
    ap.add_argument("--make_preview", action="store_true", help="生成每个病例的第1页预览图（推荐）")
    args = ap.parse_args()

    indir = Path(args.indir)
    outdir = ensure_dir(Path(args.outdir))
    prevdir = ensure_dir(outdir / "previews")

    exts = {".pdf",".docx",".doc",".png",".jpg",".jpeg",".tif",".tiff",".bmp"}
    files = [p for p in indir.rglob("*") if p.is_file() and p.suffix.lower() in exts]
    files.sort()

    rows = []
    index_rows = []

    for p in files:
        sample_id = guess_sample_id_from_name(p.name)
        ftype = p.suffix.lower().lstrip(".")
        preview_path = ""

        # ---- preview ----
        if args.make_preview:
            out_png = prevdir / f"{sample_id}__{p.stem[:30]}.png"
            ok = False
            if p.suffix.lower() == ".pdf":
                ok = render_pdf_preview(p, out_png)
            elif p.suffix.lower() in [".png",".jpg",".jpeg",".tif",".tiff",".bmp"]:
                ok = image_to_preview(p, out_png)
            elif p.suffix.lower() in [".docx",".doc"]:
                # docx/doc 常常是扫描图，预览不稳定；先不强做
                ok = False

            if ok:
                preview_path = str(out_png)

        # ---- text extraction (only if possible) ----
        text = ""
        has_text = False
        extracted = {}

        if p.suffix.lower() == ".pdf":
            has_text = pdf_has_extractable_text(p)
            if has_text:
                text = extract_text_from_pdf(p, max_pages=args.max_pages_text)
                extracted = extract_fields_from_text(text)

        elif p.suffix.lower() == ".docx":
            text = extract_text_from_docx(p)
            has_text = len(text.strip()) >= 50
            if has_text:
                extracted = extract_fields_from_text(text)

        # build excel row (auto-fill only when confident)
        row = {
            "SampleID": sample_id,
            "Group(HS/ZS/NC)": "",   # 需要你们人工填，或你们文件名有HS/ZS/NC可自行规则填
            "SourceFile": str(p),
            "PreviewLink": preview_path,
            "SamplingTime": extracted.get("SamplingTime_parsed") or extracted.get("SamplingTime",""),
            "OnsetTime": extracted.get("OnsetTime_parsed") or extracted.get("OnsetTime",""),
            "FirstArrivalTime": extracted.get("FirstArrivalTime_parsed") or extracted.get("FirstArrivalTime",""),
            "OnsetHours": extracted.get("OnsetHours",""),
            "NIHSS": extracted.get("NIHSS",""),
            "Age": extracted.get("Age",""),
            "Sex(M/F)": extracted.get("Sex",""),
            "HTN(Yes/No)": extracted.get("HTN",""),
            "DM(Yes/No)": extracted.get("DM",""),
            "Meds": extracted.get("Meds_raw",""),
            "Thrombectomy(Yes/No)": extracted.get("Thrombectomy",""),
            "Thrombolysis(Yes/No)": extracted.get("Thrombolysis",""),
            "Notes": ""
        }
        rows.append(row)

        index_rows.append({
            "sample_id": sample_id,
            "file": str(p),
            "type": ftype,
            "has_extractable_text": bool(has_text),
            "preview": preview_path,
            "extracted_json": json.dumps(extracted, ensure_ascii=False)
        })

    # save index
    pd.DataFrame(index_rows).to_csv(outdir / "case_index.csv", index=False, encoding="utf-8-sig")

    # make excel template
    out_xlsx = outdir / "ClinicalEntry_Template.xlsx"
    make_excel_template(rows, out_xlsx)

    print("✅ Done.")
    print(f"Found files: {len(files)}")
    print(f"Index CSV : {outdir / 'case_index.csv'}")
    print(f"Template  : {out_xlsx}")
    if args.make_preview:
        print(f"Previews  : {prevdir}")

if __name__ == "__main__":
    main()
